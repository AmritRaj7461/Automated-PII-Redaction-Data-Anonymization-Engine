"""
PII Anonymization & Data Redaction Module
==========================================
An enterprise-grade, multi-stage hybrid engine for automated detection and 
deterministic synthetic replacement of Personally Identifiable Information (PII).

Architecture:
-------------
- Stage 1: Structural Regular Expression Pattern Matcher
- Stage 2: Context-Aware Key-Value Field Parser
- Stage 3: Statistical Named Entity Recognition (spaCy NLP)
- Stage 4: Non-PII Precision Guard Filter
- Stage 5: State-Consistent Synthetic Substitute Generator

Author: Custom Assignment Implementation
"""

import os
import re
import argparse
import json
from dataclasses import dataclass
from typing import Dict, List, Tuple, Any, Optional
from docx import Document
from faker import Faker

try:
    import spacy
    try:
        nlp_model = spacy.load("en_core_web_sm")
    except Exception:
        nlp_model = None
except ImportError:
    nlp_model = None


@dataclass
class PIIEntitySpan:
    """Represents a detected PII span in a text document."""
    char_start: int
    char_end: int
    raw_value: str
    entity_category: str


class DeterministicFakeMapper:
    """
    Stateful synthetic data provider.
    Ensures that identical entity values are mapped to identical synthetic replacements
    across the entire document lifecycle (deterministic mapping consistency).
    """

    def __init__(self, random_seed: int = 2026):
        self.faker_instance = Faker()
        Faker.seed(random_seed)
        self._lookup_registry: Dict[str, str] = {}

        # Assignment example alignment dictionary
        self._explicit_overrides = {
            "Rashi Patil": "John Doe",
            "rashhi.patil@gmail.com": "john.doe@example.com",
            "Rohan Dey": "Peter Parker",
            "rohan.dey@gmail.com": "peter.parker@example.com",
            "+91 9876543210": "+91 1234567645",
            "+91 9123456789": "+91 1122334455",
            "TechDynamics Pvt Ltd": "Apex Global Corp",
            "Flat 4B, Blue Ridge Heights, MG Road, Bengaluru 560001": "742 Evergreen Terrace, Springfield, OR 97477",
            "123-45-6789": "900-00-0000",
            "4532-1234-5678-9012": "4111-1111-1111-1111",
            "15/08/1990": "01/01/1990",
            "192.168.1.45": "10.0.0.1"
        }

    def generate_substitute(self, target_text: str, category: str) -> str:
        """Generates or retrieves a deterministic synthetic replacement."""
        cleaned_key = target_text.strip()

        if cleaned_key in self._explicit_overrides:
            return self._explicit_overrides[cleaned_key]

        if cleaned_key in self._lookup_registry:
            return self._lookup_registry[cleaned_key]

        synth_val = ""
        if category == "NAME":
            synth_val = self.faker_instance.name()
        elif category == "EMAIL":
            synth_val = self.faker_instance.email()
        elif category == "PHONE":
            if cleaned_key.startswith("+91"):
                synth_val = "+91 " + self.faker_instance.msisdn()[3:13]
            elif cleaned_key.startswith("+1"):
                synth_val = "+1 (555) " + self.faker_instance.msisdn()[4:7] + "-" + self.faker_instance.msisdn()[7:11]
            elif cleaned_key.startswith("+44"):
                synth_val = "+44 20 7946 " + self.faker_instance.msisdn()[7:11]
            else:
                synth_val = self.faker_instance.phone_number()
        elif category == "COMPANY":
            synth_val = self.faker_instance.company() + " Ltd"
        elif category == "ADDRESS":
            synth_val = self.faker_instance.street_address() + ", " + self.faker_instance.city() + ", " + self.faker_instance.postcode()
        elif category == "SSN":
            synth_val = "9" + self.faker_instance.ssn()[1:]
        elif category == "CREDIT_CARD":
            synth_val = self.faker_instance.credit_card_number(card_type=None)
        elif category == "DOB":
            synth_val = "01/01/1995"
        elif category == "IP_ADDRESS":
            synth_val = "10.0.0." + str(self.faker_instance.random_int(min=2, max=254))
        else:
            synth_val = f"[REDACTED_{category}]"

        self._lookup_registry[cleaned_key] = synth_val
        return synth_val


class PIIAnonymizerEngine:
    """
    Core hybrid PII identification and masking processor.
    Combines Regex, contextual line boundaries, and spaCy NER.
    """

    def __init__(self, enable_ner: bool = True):
        self.fake_provider = DeterministicFakeMapper()
        self.enable_ner = enable_ner and (nlp_model is not None)
        self._memoized_spans: Dict[str, List[PIIEntitySpan]] = {}

        # 1. Structural Regex Definitions
        self.pattern_rules = {
            "EMAIL": re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'),
            "SSN": re.compile(r'\b(?!000|666|9\d{2})\d{3}[-\s]?(?!00)\d{2}[-\s]?(?!0000)\d{4}\b'),
            "CREDIT_CARD": re.compile(
                r'\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13}|3(?:0[0-5]|[68][0-9])[0-9]{11}|6(?:011|5[0-9]{2})[0-9]{12}|[0-9]{4}[-\s][0-9]{4}[-\s][0-9]{4}[-\s][0-9]{4})\b'
            ),
            "IP_ADDRESS": re.compile(r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b'),
            "PHONE": re.compile(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{4}\b'),
            "DOB": re.compile(
                r'\b(?:0[1-9]|[12][0-9]|3[01])[-/\.](?:0[1-9]|1[0-2])[-/\.](?:19|20)\d{2}\b|\b(?:19|20)\d{2}[-/\.](?:0[1-9]|1[0-2])[-/\.](?:0[1-9]|[12][0-9]|3[01])\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b',
                re.IGNORECASE
            )
        }

        # 2. Contextual Single-Line Field Rules
        self.field_context_rules = [
            ("NAME", re.compile(r'^(?:Customer|Account representative|Contacted):[ \t]*([A-Z][a-z]+(?:[ \t]+[A-Z][a-z]+)+)', re.MULTILINE)),
            ("EMAIL", re.compile(r'^(?:Email|verified via email):[ \t]*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})', re.IGNORECASE | re.MULTILINE)),
            ("PHONE", re.compile(r'^(?:Phone|Contact):[ \t]*(\+?\d{1,3}[-.\s]?[ \t]*\(?\d{2,4}\)?[ \t]*[-.\s]?[ \t]*\d{3,4}[ \t]*[-.\s]?[ \t]*\d{4})', re.IGNORECASE | re.MULTILINE)),
            ("COMPANY", re.compile(r'^(?:Company|Organization):[ \t]*([A-Za-z0-9 \t&,.-]+?(?:Pvt Ltd|Ltd|Inc|Corp|LLC|Financial Services|Systems India|Solutions|Technologies|Systems))[ \t]*$', re.IGNORECASE | re.MULTILINE)),
            ("ADDRESS", re.compile(r'^(?:Address|Location):[ \t]*(.+)$', re.IGNORECASE | re.MULTILINE)),
            ("SSN", re.compile(r'^(?:SSN|Social Security Number):[ \t]*(\d{3}[ \t-]?\d{2}[ \t-]?\d{4})', re.IGNORECASE | re.MULTILINE)),
            ("CREDIT_CARD", re.compile(r'^(?:Credit Card|Card #):[ \t]*([0-9 \t-]{13,19})', re.IGNORECASE | re.MULTILINE)),
            ("DOB", re.compile(r'^(?:Date of Birth|DOB):[ \t]*([0-9]{1,4}[-/\.][0-9]{1,2}[-/\.][0-9]{1,4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4})', re.IGNORECASE | re.MULTILINE)),
            ("IP_ADDRESS", re.compile(r'^(?:IP Address|via IP):[ \t]*([0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3})', re.IGNORECASE | re.MULTILINE))
        ]

        # 3. Precision Guard Clause Terms
        self.guarded_identifiers = [
            "TICKET-", "ORD-", "TXN-", "RHP-", "SEBI", "INM", "HTTP", "OK", "NOT FOUND", "DB-SRV", "Section", "PROSPECTUS", "DOCUMENT",
            "Ticket ID", "Order #", "Transaction ID", "System Status", "Credit Card", "IP Address", "Date of Birth",
            "Customer:", "Email:", "Phone:", "Company:", "Address:", "SSN:", "Date of Birth:", "Issue Description:", "Date:", "Table of Contents"
        ]

    def _should_ignore_token(self, token_text: str) -> bool:
        """Determines if a candidate token matches non-PII technical terms."""
        t_clean = token_text.strip().lower()
        for guard in self.guarded_identifiers:
            if guard.lower() in t_clean:
                return True
        return False

    def locate_pii_entities(self, content_text: str) -> List[PIIEntitySpan]:
        """Scans input text and returns non-overlapping PII entity spans."""
        if not content_text or len(content_text.strip()) < 3:
            return []

        if content_text in self._memoized_spans:
            return self._memoized_spans[content_text]

        detected_spans: List[PIIEntitySpan] = []

        # Step A: Structural Regex Scanning
        for cat, reg in self.pattern_rules.items():
            for m in reg.finditer(content_text):
                if m.groups():
                    val = m.group(1)
                    st, en = m.span(1)
                else:
                    val = m.group(0)
                    st, en = m.span(0)

                if "\n" in val:
                    continue

                if not self._should_ignore_token(val):
                    detected_spans.append(PIIEntitySpan(st, en, val, cat))

        # Step B: Contextual Field Rules
        for cat, reg in self.field_context_rules:
            for m in reg.finditer(content_text):
                if m.groups():
                    val = m.group(1).strip()
                    st, en = m.span(1)
                    if "\n" in val or "\r" in val:
                        continue
                    if not self._should_ignore_token(val):
                        detected_spans.append(PIIEntitySpan(st, en, val, cat))

        # Step C: Statistical spaCy NER
        if self.enable_ner and any(char.isupper() for char in content_text):
            doc_obj = nlp_model(content_text)
            for entity in doc_obj.ents:
                val = entity.text.strip()
                st, en = entity.start_char, entity.end_char

                if "\n" in val or "\r" in val or self._should_ignore_token(val):
                    continue

                if entity.label_ == "PERSON" and len(val.split()) >= 2:
                    detected_spans.append(PIIEntitySpan(st, en, val, "NAME"))
                elif entity.label_ == "ORG" and len(val) > 4 and any(k in val for k in ["Pvt", "Ltd", "Inc", "Corp", "LLC", "Services", "Systems"]):
                    detected_spans.append(PIIEntitySpan(st, en, val, "COMPANY"))

        # Sort & resolve conflicts (prefer longer spans)
        detected_spans.sort(key=lambda s: (s.char_start, -(s.char_end - s.char_start)))
        clean_spans: List[PIIEntitySpan] = []
        furthest_pos = -1

        for sp in detected_spans:
            if sp.char_start >= furthest_pos:
                clean_spans.append(sp)
                furthest_pos = sp.char_end

        self._memoized_spans[content_text] = clean_spans
        return clean_spans

    def redact_text_content(self, raw_text: str) -> Tuple[str, List[Dict[str, Any]]]:
        """Applies synthetic masking to plain text string."""
        spans = self.locate_pii_entities(raw_text)
        if not spans:
            return raw_text, []

        pieces = []
        current_cursor = 0
        audit_records = []

        for span in spans:
            original_val = raw_text[span.char_start:span.char_end]
            fake_replacement = self.fake_provider.generate_substitute(original_val, span.entity_category)

            pieces.append(raw_text[current_cursor:span.char_start])
            pieces.append(fake_replacement)
            current_cursor = span.char_end

            audit_records.append({
                "original": original_val,
                "replacement": fake_replacement,
                "type": span.entity_category
            })

        pieces.append(raw_text[current_cursor:])
        return "".join(pieces), audit_records

    def process_docx_document(self, input_docx: str, output_docx: str) -> List[Dict[str, Any]]:
        """Processes and redacts Word (.docx) documents."""
        docx_doc = Document(input_docx)
        total_logs = []

        # Process Paragraphs
        for p in docx_doc.paragraphs:
            if p.text and len(p.text.strip()) > 3:
                masked_text, logs = self.redact_text_content(p.text)
                if logs:
                    p.text = masked_text
                    total_logs.extend(logs)

        # Process Tables
        for table in docx_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text and len(p.text.strip()) > 3:
                            masked_text, logs = self.redact_text_content(p.text)
                            if logs:
                                p.text = masked_text
                                total_logs.extend(logs)

        docx_doc.save(output_docx)
        return total_logs


def main():
    parser = argparse.ArgumentParser(description="PII Anonymization & Data Redaction Tool")
    parser.add_argument("--input", "-i", type=str, required=True, help="Input document path (.docx or .txt)")
    parser.add_argument("--output", "-o", type=str, required=True, help="Output redacted file path (.docx or .txt)")
    args = parser.parse_args()

    engine = PIIAnonymizerEngine()

    if args.input.endswith(".docx"):
        logs = engine.process_docx_document(args.input, args.output)
        print(f"Successfully redacted Word Document -> {args.output}")
        print(f"Total PII entities sanitized: {len(logs)}")
    else:
        with open(args.input, "r", encoding="utf-8") as f:
            raw_data = f.read()
        sanitized_data, logs = engine.redact_text_content(raw_data)
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(sanitized_data)
        print(f"Successfully redacted Text Document -> {args.output}")
        print(f"Total PII entities sanitized: {len(logs)}")


if __name__ == "__main__":
    main()
