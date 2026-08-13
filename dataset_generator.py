"""
Benchmark Corpus & Ground Truth Annotations Generator
=====================================================
Constructs annotated evaluation corpora and test Word (.docx) documents
containing PII entities across 9 distinct categories and non-PII technical IDs.
"""

import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

AUDIT_CORPUS_BLOCKS = [
    {
        "meta": "RED HERRING PROSPECTUS - CUSTOMER SUPPORT TICKET LOG",
        "text": "DOCUMENT REFERENCE: RHP-2026-X99 | SEBI REGISTRATION NO: INM000010921\n"
                "CONFIDENTIAL SUPPORT TICKET AUDIT LOG\n"
                "Date: 2026-08-13 10:15:30 UTC | System Status: 200 OK | Node: DB-SRV-01\n"
    },
    {
        "ticket_id": "TICKET-94021",
        "text": "[2026-08-13 10:17:02] [LOG] Ticket ID: TICKET-94021 | Order #: ORD-2024-8819 | Transaction ID: TXN-99410283\n"
                "Customer: Rashi Patil\n"
                "Email: rashhi.patil@gmail.com\n"
                "Phone: +91 9876543210\n"
                "Company: TechDynamics Pvt Ltd\n"
                "Address: Flat 4B, Blue Ridge Heights, MG Road, Bengaluru 560001\n"
                "SSN: 123-45-6789\n"
                "Credit Card: 4532-1234-5678-9012\n"
                "Date of Birth: 15/08/1990\n"
                "IP Address: 192.168.1.45\n"
                "Issue Description: Customer requested clarification on Section 4.2 of the Red Herring Prospectus. "
                "The payment of $15,000,000 for equity subscription was processed under TXN-99410283.",
        "pii": [
            {"type": "NAME", "value": "Rashi Patil"},
            {"type": "EMAIL", "value": "rashhi.patil@gmail.com"},
            {"type": "PHONE", "value": "+91 9876543210"},
            {"type": "COMPANY", "value": "TechDynamics Pvt Ltd"},
            {"type": "ADDRESS", "value": "Flat 4B, Blue Ridge Heights, MG Road, Bengaluru 560001"},
            {"type": "SSN", "value": "123-45-6789"},
            {"type": "CREDIT_CARD", "value": "4532-1234-5678-9012"},
            {"type": "DOB", "value": "15/08/1990"},
            {"type": "IP_ADDRESS", "value": "192.168.1.45"}
        ],
        "non_pii": ["TICKET-94021", "ORD-2024-8819", "TXN-99410283", "$15,000,000", "Section 4.2"]
    },
    {
        "ticket_id": "TICKET-88120",
        "text": "[2026-08-13 11:04:15] [LOG] Ticket ID: TICKET-88120 | Order #: ORD-2026-9901 | Transaction ID: TXN-55410928\n"
                "Customer: Rohan Dey\n"
                "Email: rohan.dey@gmail.com\n"
                "Phone: +91 9123456789\n"
                "Company: Acme Financial Services\n"
                "Address: 124 Park Avenue, Suite 300, New York, NY 10001\n"
                "SSN: 987-65-4321\n"
                "Credit Card: 5412-7512-3412-3456\n"
                "Date of Birth: 1985-11-23\n"
                "IP Address: 172.16.254.1\n"
                "Issue Description: Customer reported HTTP 404 NOT FOUND error when submitting institutional bid via IP 172.16.254.1. "
                "Account representative Rohan Dey confirmed retry.",
        "pii": [
            {"type": "NAME", "value": "Rohan Dey"},
            {"type": "EMAIL", "value": "rohan.dey@gmail.com"},
            {"type": "PHONE", "value": "+91 9123456789"},
            {"type": "COMPANY", "value": "Acme Financial Services"},
            {"type": "ADDRESS", "value": "124 Park Avenue, Suite 300, New York, NY 10001"},
            {"type": "SSN", "value": "987-65-4321"},
            {"type": "CREDIT_CARD", "value": "5412-7512-3412-3456"},
            {"type": "DOB", "value": "1985-11-23"},
            {"type": "IP_ADDRESS", "value": "172.16.254.1"}
        ],
        "non_pii": ["TICKET-88120", "ORD-2026-9901", "TXN-55410928", "HTTP 404 NOT FOUND"]
    },
    {
        "ticket_id": "TICKET-10492",
        "text": "[2026-08-13 12:30:45] [LOG] Ticket ID: TICKET-10492 | Order #: ORD-2026-7742 | Transaction ID: TXN-33210984\n"
                "Customer: Sarah Jenkins\n"
                "Email: sarah.j@acme-corp.org\n"
                "Phone: +1 (555) 234-5678\n"
                "Company: Cyberdyne Systems India\n"
                "Address: Plot 15, Sector 18, Cyber City, Gurgaon 122002\n"
                "SSN: 456-78-1234\n"
                "Credit Card: 3782-822456-91001\n"
                "Date of Birth: March 12, 1994\n"
                "IP Address: 10.0.4.12\n"
                "Issue Description: Inquiry regarding prospectus share allocation rules. Contacted Sarah Jenkins at Cyberdyne Systems India.",
        "pii": [
            {"type": "NAME", "value": "Sarah Jenkins"},
            {"type": "EMAIL", "value": "sarah.j@acme-corp.org"},
            {"type": "PHONE", "value": "+1 (555) 234-5678"},
            {"type": "COMPANY", "value": "Cyberdyne Systems India"},
            {"type": "ADDRESS", "value": "Plot 15, Sector 18, Cyber City, Gurgaon 122002"},
            {"type": "SSN", "value": "456-78-1234"},
            {"type": "CREDIT_CARD", "value": "3782-822456-91001"},
            {"type": "DOB", "value": "March 12, 1994"},
            {"type": "IP_ADDRESS", "value": "10.0.4.12"}
        ],
        "non_pii": ["TICKET-10492", "ORD-2026-7742", "TXN-33210984"]
    },
    {
        "ticket_id": "TICKET-77291",
        "text": "[2026-08-13 14:15:10] [LOG] Ticket ID: TICKET-77291 | Order #: ORD-2026-4410 | Transaction ID: TXN-11029384\n"
                "Customer: Vikram Malhotra\n"
                "Email: v.malhotra@techdynamics.in\n"
                "Phone: +44 20 7946 0912\n"
                "Company: Global Wealth Solutions\n"
                "Address: 45 Oxford Street, Westminster, London W1D 2DZ\n"
                "SSN: 789-01-2345\n"
                "Credit Card: 6011-1111-2222-3333\n"
                "Date of Birth: 04-07-1978\n"
                "IP Address: 203.0.113.195\n"
                "Issue Description: Vikram Malhotra from Global Wealth Solutions submitted escrow deposit. Verified via email v.malhotra@techdynamics.in.",
        "pii": [
            {"type": "NAME", "value": "Vikram Malhotra"},
            {"type": "EMAIL", "value": "v.malhotra@techdynamics.in"},
            {"type": "PHONE", "value": "+44 20 7946 0912"},
            {"type": "COMPANY", "value": "Global Wealth Solutions"},
            {"type": "ADDRESS", "value": "45 Oxford Street, Westminster, London W1D 2DZ"},
            {"type": "SSN", "value": "789-01-2345"},
            {"type": "CREDIT_CARD", "value": "6011-1111-2222-3333"},
            {"type": "DOB", "value": "04-07-1978"},
            {"type": "IP_ADDRESS", "value": "203.0.113.195"}
        ],
        "non_pii": ["TICKET-77291", "ORD-2026-4410", "TXN-11029384"]
    }
]


def generate_dataset():
    """Builds corpus, DOCX test document, and ground-truth JSON files."""
    work_dir = os.path.dirname(os.path.abspath(__file__))

    # 1. DOCX Test Document
    docx_file = os.path.join(work_dir, "input_ticket_log.docx")
    docx_obj = Document()

    title_para = docx_obj.add_paragraph()
    r_t = title_para.add_run("RED HERRING PROSPECTUS - CUSTOMER SUPPORT TICKET LOG")
    r_t.bold = True
    r_t.font.size = Pt(16)
    r_t.font.color.rgb = RGBColor(0, 51, 102)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_para = docx_obj.add_paragraph("DOCUMENT REFERENCE: RHP-2026-X99 | CONFIDENTIAL PII AUDIT DATA")
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for blk in AUDIT_CORPUS_BLOCKS:
        if "ticket_id" in blk:
            head_p = docx_obj.add_paragraph()
            r_h = head_p.add_run(f"--- Support Ticket: {blk['ticket_id']} ---")
            r_h.bold = True
            r_h.font.size = Pt(12)
            r_h.font.color.rgb = RGBColor(128, 0, 0)

            body_p = docx_obj.add_paragraph(blk["text"])
            body_p.paragraph_format.line_spacing = 1.15
            body_p.paragraph_format.space_after = Pt(10)

    docx_obj.save(docx_file)
    print(f"Constructed Benchmark DOCX File: {docx_file}")

    # 2. Ground Truth JSON Annotations
    gt_file = os.path.join(work_dir, "ground_truth.json")
    annotations = []

    for blk in AUDIT_CORPUS_BLOCKS:
        if "pii" in blk:
            for entity in blk["pii"]:
                annotations.append({
                    "ticket_id": blk["ticket_id"],
                    "type": entity["type"],
                    "value": entity["value"]
                })
        if "non_pii" in blk:
            for non_entity in blk["non_pii"]:
                annotations.append({
                    "ticket_id": blk["ticket_id"],
                    "type": "NON_PII",
                    "value": non_entity
                })

    with open(gt_file, "w", encoding="utf-8") as f:
        json.dump(annotations, f, indent=2)
    print(f"Constructed Ground Truth Annotations JSON: {gt_file}")


if __name__ == "__main__":
    generate_dataset()
